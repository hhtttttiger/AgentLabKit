import io
from docx import Document as DocxDocument
from retrieval.engines.local_engine.document_loaders.base_loader import ServiceLoader
from retrieval.engines.local_engine.document_loaders.base import Document, Page

class DocxPandocLoader(ServiceLoader):
    def load(self, file_bytes: bytes, file_info, **kwargs) -> Document:
        stream = io.BytesIO(file_bytes)
        doc = DocxDocument(stream)
        
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
            
        content = "\n".join(full_text)
        page_list = [Page(content=content, page_number=1)]
        
        return Document(content=content, page_list=page_list, is_md=False)
